"""
Script to generate SharePoint Analysis Report as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('DigiHub SharePoint Document Analysis Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Date: February 5, 2026')
doc.add_paragraph('Purpose: Analysis of SharePoint documents for DigiHub Chatbot knowledge base optimization')
doc.add_paragraph()

# Section 1: Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report analyzes the SharePoint document library used as the knowledge source for the DigiHub Chatbot. '
    'The analysis identifies gaps between available documents and indexed content in CosmosDB, '
    'and provides recommendations for improving chatbot response quality for generic queries.'
)

# Section 2: Total SharePoint Overview
doc.add_heading('2. Total SharePoint Overview (All Files)', level=1)
doc.add_paragraph('This section covers ALL files in SharePoint, including those not indexed for the chatbot.')

table1 = doc.add_table(rows=13, cols=4)
table1.style = 'Table Grid'
# Header
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Folder'
hdr_cells[1].text = 'Total Size'
hdr_cells[2].text = 'Total Files'
hdr_cells[3].text = 'Notes'

data1 = [
    ('WorldTracer', '4.3 GB', '538', 'Largest by size'),
    ('Euro Customer Advisory Board', '1.4 GB', '68', 'Meeting notes'),
    ('Bag Manager', '239 MB', '378', 'User guides, release notes, FAQs'),
    ('Billing', '27 MB', '64', 'User guides, billing updates'),
    ('SITA Messaging', '10 MB', '24', 'Documentation'),
    ('Airport Committee', '4.1 MB', '11', 'Meeting minutes'),
    ('Operational Support', '3.4 MB', '2', 'CI Analysis, Incidents, AIR Dashboard'),
    ('Airport Management Solutions', '1.0 MB', '8', 'Test files'),
    ('General Info', '192 KB', '2', 'Main DigiHub guide, Acronym list'),
    ('SITA AeroPerformance', '52 KB', '3', 'Documentation'),
    ('APAC Customer Advisory Board', '28 KB', '3', 'Meeting docs'),
    ('Airport Solutions', '0 B', '0', 'Empty folder'),
]

for i, row_data in enumerate(data1):
    row = table1.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Total: 6.0 GB, ~1,100 files')

# File types breakdown
doc.add_heading('2.1 File Types Breakdown', level=2)
table_ft = doc.add_table(rows=7, cols=2)
table_ft.style = 'Table Grid'
table_ft.rows[0].cells[0].text = 'File Type'
table_ft.rows[0].cells[1].text = 'Count'
ft_data = [('PDF', '678'), ('DOCX', '207'), ('PPTX', '78'), ('MP4 (video)', '48'), ('XLSX (Excel)', '32'), ('Other', '57')]
for i, (ft, cnt) in enumerate(ft_data):
    table_ft.rows[i + 1].cells[0].text = ft
    table_ft.rows[i + 1].cells[1].text = cnt

doc.add_paragraph()

# Section 3: Indexed Files Only
doc.add_heading('3. Indexed Files for Chatbot (Excluding Videos & Excel)', level=1)
doc.add_paragraph(
    'The chatbot indexes only text-based documents: PDF, DOCX, DOC, PPTX, PPT, TXT. '
    'Videos (MP4) and Excel files (XLSX) are NOT indexed.'
)

table2 = doc.add_table(rows=13, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Folder'
hdr2[1].text = 'Indexed Size'
hdr2[2].text = 'Indexed Files'
hdr2[3].text = 'Excluded Files'

data2 = [
    ('WorldTracer', '1.0 GB', '463', '75 (videos/excel)'),
    ('Bag Manager', '204 MB', '364', '14'),
    ('Euro Customer Advisory Board', '198 MB', '64', '4'),
    ('Billing', '18 MB', '61', '3'),
    ('SITA Messaging', '10 MB', '24', '0'),
    ('Airport Committee', '4.1 MB', '11', '0'),
    ('Operational Support', '3.4 MB', '2', '0'),
    ('Airport Management Solutions', '964 KB', '5', '3'),
    ('General Info', '192 KB', '2', '0'),
    ('SITA AeroPerformance', '36 KB', '2', '1'),
    ('APAC Customer Advisory Board', '4 KB', '1', '2'),
    ('Airport Solutions', '0', '0', '0'),
]

for i, row_data in enumerate(data2):
    row = table2.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Total Indexed Files: 999 files')
doc.add_paragraph('Total Excluded Files: 102 files (48 videos, 32 Excel, 12 images, 10 other)')

# Section 4: DB Chunks vs Indexed Files
doc.add_heading('4. CosmosDB Chunks vs SharePoint Files (Gap Analysis)', level=1)
doc.add_paragraph(
    'This comparison reveals significant gaps between indexed files and database chunks. '
    'A large number of indexed files means documents exist but may not be properly chunked/ingested.'
)

table3 = doc.add_table(rows=11, cols=4)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Service Line'
hdr3[1].text = 'Service ID'
hdr3[2].text = 'DB Chunks'
hdr3[3].text = 'Indexed Files'

data3 = [
    ('WorldTracer', '(various)', '2,968', '463'),
    ('Bag Manager', '(various)', '2,868', '364'),
    ('Euro Customer Advisory Board', '(various)', '711', '64'),
    ('(No serviceName)', 'NULL', '519', 'N/A'),
    ('SITA Messaging', '(various)', '120', '24'),
    ('Billing', '400', '103', '61'),
    ('Airport Committee', '(various)', '45', '11'),
    ('General Info', '0', '10', '2'),
    ('Operational Support', '440', '10', '2'),
    ('SITA AeroPerformance', '(various)', '2', '2'),
]

for i, row_data in enumerate(data3):
    row = table3.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# Critical findings
doc.add_heading('4.1 Critical Findings (Post-Indexing - Feb 5, 2026)', level=2)
findings = doc.add_paragraph()
findings.add_run('WorldTracer: ').bold = True
findings.add_run('Indexing COMPLETE. 2,968 chunks from 463 files. Now the largest service line in the database.\n\n')
findings.add_run('General Info: ').bold = True
findings.add_run('Only 10 chunks from 2 files. Contains main DigiHub user guide - critical for generic queries.\n\n')
findings.add_run('Operational Support: ').bold = True
findings.add_run('Only 10 chunks from 2 files. Contains CI Analysis, Incidents documentation.\n\n')
findings.add_run('Euro CAB Dominance: ').bold = True
findings.add_run('711 chunks from meeting notes may dilute search results for generic queries.\n\n')
findings.add_run('Total Chunks: ').bold = True
findings.add_run('7,356 chunks across all service lines.')

# Section 5: Key Documents Content
doc.add_heading('5. Key User Guide Documents (Content Analysis)', level=1)

doc.add_heading('5.1 General Info Folder (serviceNameid = 0)', level=2)
doc.add_paragraph('Files: 2 | Chunks in DB: 10 | Content Type: DigiHubUserGuide')

doc.add_heading('Digihub (v2026-02-03).docx', level=3)
content1 = doc.add_paragraph()
content1.add_run('Main DigiHub user guide containing:\n')
bullets1 = [
    'What is DigiHub / What can I do with DigiHub',
    'Billing & Payment features overview',
    'Operational Support overview',
    'Airport Solutions',
    'SITA Connect Go / SDN',
    'Product Knowledge Base (Messaging, WorldTracer, Bag Manager)',
    'DigiHub Library',
    'Idea Hub / Learning Hub',
    'AI Assistant chatbot language support',
    'How to create support tickets',
    'How to update user profile'
]
for b in bullets1:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Acronym List final.pdf', level=3)
doc.add_paragraph('Comprehensive acronym guide with aviation/IT industry terms (AAC, ACARS, API, ARINC, etc.)')

doc.add_heading('5.2 Operational Support Folder (serviceNameid = 440)', level=2)
doc.add_paragraph('Files: 2 | Chunks in DB: 10 | Content Type: UserGuide')

doc.add_heading('DigiHub - AIR Dashboard - User guide (v1.1).pdf', level=3)
bullets2 = [
    'CI Analysis - configuration items, top issues, resolution codes',
    'Active Incidents - status, priority, age groups',
    'Smart View Page - incidents by module, trends, SLA compliance',
    'GEO Analysis - airport/airline incident distribution',
    'Glossary - FCR, MTTR, SLA definitions',
    'Troubleshooting tips'
]
for b in bullets2:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Operational Support - User Guide.pdf', level=3)
doc.add_paragraph('Guide for raising and tracking incidents and service requests.')

doc.add_heading('5.3 Billing Folder (serviceNameid = 400)', level=2)
doc.add_paragraph('Files: 61 | Chunks in DB: 103 | Content Types: DigiHubUserGuide, BillingUpdate')
bullets3 = [
    'Billing & Payment User Guide - subscribe, navigate, search invoices, download',
    'E-Invoicing rules management',
    'Billing support cases',
    'Monthly billing updates (2022-2025)'
]
for b in bullets3:
    doc.add_paragraph(b, style='List Bullet')

# Section 6: Problem Analysis
doc.add_heading('6. Generic Query Problem Analysis', level=1)

doc.add_heading('6.1 The Problem', level=2)
doc.add_paragraph(
    'Generic queries like "How do I export to CSV?", "What is CI Analysis?", "How to search for incidents?" '
    'return out-of-scope responses instead of finding answers in DigiHub documentation.'
)

doc.add_heading('6.2 Root Cause', level=2)
causes = [
    'No service line detected from generic query',
    'System searches ALL service lines (7,356 chunks post-indexing)',
    'Euro Customer Advisory Board has 711 meeting note chunks that dilute results',
    'Vector search returns semantically similar but irrelevant content',
    'Relevance judge filters out everything, returning out-of-scope'
]
for c in causes:
    doc.add_paragraph(c, style='List Number')

doc.add_heading('6.3 Evidence', level=2)
doc.add_paragraph(
    'CosmosDB query for "csv", "export", "ci analysis", "incidents" in chunk field returned EMPTY, '
    'yet the headings "CI Analysis" and "View Company Incidents" exist in Operational Support chunks.'
)

# Section 7: Recommendations
doc.add_heading('7. Recommendations', level=1)

doc.add_heading('7.1 Immediate Fix: Default Service Lines for Generic Queries', level=2)
doc.add_paragraph(
    'When no service line is detected from the query, default to searching:\n'
    '- serviceNameid = 0 (General Info)\n'
    '- serviceNameid = 440 (Operational Support)\n\n'
    'This ensures generic DigiHub questions find user guide content instead of meeting notes.'
)

doc.add_heading('7.2 Data Quality Improvements', level=2)
improvements = [
    'WorldTracer: Indexing in progress - monitor completion',
    'General Info: Consider chunking the main DigiHub guide into more granular chunks',
    'Content Type Tagging: Ensure all user guides have contentType = "UserGuide" or "DigiHubUserGuide"',
    'Review 519 chunks with NULL serviceName - assign appropriate service lines'
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_heading('7.3 Search Quality Improvements', level=2)
search_imp = [
    'Boost user guide content types in search ranking',
    'Implement two-stage retrieval: search user guides first, then expand',
    'Consider reducing weight of meeting notes in generic queries'
]
for s in search_imp:
    doc.add_paragraph(s, style='List Bullet')

# Section 8: Service Line Reference
doc.add_heading('8. Service Line ID Reference', level=1)
table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Service Line'
table4.rows[0].cells[1].text = 'ID'
table4.rows[0].cells[2].text = 'Content Types'

ref_data = [
    ('General Info', '0', 'DigiHubUserGuide'),
    ('Billing', '400', 'DigiHubUserGuide, BillingUpdate'),
    ('Operational Support', '440', 'UserGuide'),
    ('(No serviceName)', 'NULL', 'Various'),
    ('Others', 'Various', 'MeetingNotes, Others, Survey'),
]
for i, row_data in enumerate(ref_data):
    row = table4.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

# Section 9: Implemented Fixes
doc.add_heading('9. Implemented Fixes (Feb 5, 2026)', level=1)

doc.add_heading('9.1 Default Service Lines for Generic Queries', level=2)
doc.add_paragraph(
    'File: src/chatbot/response_generator.py\n'
    'Beads Issue: digihub-chatbot-inb (CLOSED)\n\n'
    'When no service line is detected from the query, the system now defaults to:\n'
    '- serviceNameid = 0 (General Info)\n'
    '- serviceNameid = 440 (Operational Support)\n\n'
    'This ensures generic DigiHub questions search user guide content first.'
)

doc.add_heading('9.2 Content Type Boosting for User Guides', level=2)
doc.add_paragraph(
    'File: src/services/retrieval_service.py\n'
    'Beads Issue: digihub-chatbot-tgx (CLOSED)\n\n'
    'User guide content types receive a 5% score boost:\n'
    '- UserGuide\n'
    '- DigiHubUserGuide\n\n'
    'This helps prioritize user guides over newsletters (BillingUpdate) when both are present in search results.'
)

# Section 10: Remaining Issues
doc.add_heading('10. Remaining Issues', level=1)

doc.add_heading('10.1 CI Analysis Questions Not Answered', level=2)
doc.add_paragraph('Status: DATA QUALITY ISSUE - Requires question regeneration')

ci_issue = doc.add_paragraph()
ci_issue.add_run('Problem: ').bold = True
ci_issue.add_run('"What is CI Analysis?" queries return out-of-scope responses.\n\n')
ci_issue.add_run('Root Cause: ').bold = True
ci_issue.add_run('The CI Analysis chunk exists with correct metadata, but its generated questions are:\n')

ci_questions = [
    '"How do I access the CI analysis section in DigiHub?"',
    '"How do I view configuration items and related data in DigiHub?"',
    '"How can I analyze top issues, resolution codes, and other metrics in DigiHub?"'
]
for q in ci_questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph(
    'Notice: All questions are "How do I..." format. There is NO "What is CI Analysis?" question.\n'
    'The vector search uses question embeddings, so "What is CI Analysis?" matches poorly.\n\n'
    'Evidence: The query "What is CI Analysis?" matches "What is Operational Support?" better (0.7728) '
    'than any CI Analysis-related question (0.7609).'
)

doc.add_heading('10.1.1 Solution Options', level=3)
solutions = [
    'Option A: Re-run question generation to include "What is X?" questions for each heading',
    'Option B: Add heading-based boosting (exact heading matches get score boost)',
    'Option C: Implement query expansion (prepend "How do I use" to "What is X" queries)',
    'Option D: Add chunk content to embedding alongside questions (hybrid chunk+question embedding)'
]
for s in solutions:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('10.2 Billing User Guide vs Newsletter Volume Imbalance', level=2)
doc.add_paragraph(
    'Status: PARTIALLY MITIGATED by content type boosting\n\n'
    'The Billing service line has:\n'
    '- 96 BillingUpdate chunks (newsletters)\n'
    '- 5 DigiHubUserGuide chunks (user guides)\n\n'
    'The 5% boost helps, but the 19:1 ratio means newsletters may still dominate for some queries.\n'
    'Consider higher boost or two-stage retrieval for better results.'
)

# Section 11: Test Questions
doc.add_heading('11. Recommended Test Questions', level=1)
doc.add_paragraph('Use these questions to validate chatbot response quality:')

test_table = doc.add_table(rows=9, cols=3)
test_table.style = 'Table Grid'
test_table.rows[0].cells[0].text = 'Question'
test_table.rows[0].cells[1].text = 'Expected Source'
test_table.rows[0].cells[2].text = 'Tests'

test_data = [
    ('What is DigiHub?', 'General Info', 'Generic query default service lines'),
    ('How do I create a support ticket?', 'General Info / Operational Support', 'Generic query retrieval'),
    ('How do I download an invoice?', 'Billing User Guide', 'Content type boosting (guide vs newsletter)'),
    ('What billing features are available?', 'Billing User Guide', 'Content type boosting'),
    ('How do I access CI Analysis?', 'Operational Support', 'CI Analysis "How do I" format'),
    ('What is CI Analysis?', 'Operational Support', 'CI Analysis "What is" format (MAY FAIL)'),
    ('How do I track incidents?', 'Operational Support', 'Generic query with Operational Support'),
    ('What is WorldTracer?', 'WorldTracer', 'New WorldTracer chunks working'),
]

for i, row_data in enumerate(test_data):
    row = test_table.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Note: "What is CI Analysis?" may still fail due to question embedding mismatch (see Section 10.1).')

# Save
doc.save('/Users/zohaibtanwir/projects/digihub-chatbot/digihub-chatbot-query-pipeline/SharePoint_Analysis_Report.docx')
print("Report saved to SharePoint_Analysis_Report.docx")
