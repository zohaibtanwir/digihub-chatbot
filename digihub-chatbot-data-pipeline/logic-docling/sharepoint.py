import pandas as pd
from docx import Document

from src.services.auth_service import GraphAuthService
from src.utils.indexing_utils import get_iso_datetime_to_unix_timestamp
from src.utils.logger import logger
import requests
from src.models.document_payload import GenerateSummaryPayload, Item
import asyncio
from docling.document_converter import DocumentStream
import httpx
import io
from urllib.parse import urlparse,unquote,parse_qs, quote
from src.enums.graph_urls import GraphUrls


class SharePointService:
    site_mapping = {
        "ProposalCentral": {
            "site_id": "sita365.sharepoint.com,34f77b2d-7d6d-48fa-a267-bb686dba07a9,bb486867-f245-490e-88ed-343d10890aa3",
            "doc_type": "APAC Region"
        },
        "ASLEUROpportunities": {
            "site_id": "sita365.sharepoint.com,3615dde3-b1f0-4129-b37c-00f0871c11c3,fb769df8-cdd4-4673-a714-56e429a451b9",
            "doc_type": "EURO Region"
        }
    }

    def __init__(self):
        self.auth_service = GraphAuthService()
        self.drive_id = None

    def _get_headers(self):
        headers = {
            "Authorization": f"Bearer {self.auth_service.get_access_token()}",
            'Content-Type': 'application/json'
        }
        return headers

    async def get_file_contents(self, documents: GenerateSummaryPayload):
        tasks = [
            self.download_file(
                site_id=documents.siteId,
                drive_id=documents.driveId,
                item_data=file
            )
            for file in documents.items
        ]
        contents = await asyncio.gather(*tasks)
        logger.info(f"[get_file_contents] No of Files fetched:{len(contents)}")
        return contents

    async def get_contents_sync(self, documents: GenerateSummaryPayload):

        return await self.get_file_contents(documents)

    def convert_excel_to_word_streams(self,excel_bytes,file_id):
        # Load the Excel file from bytes
        excel_io = io.BytesIO(excel_bytes)
        xls = pd.ExcelFile(excel_io, engine='openpyxl')

        word_streams = []

        # Iterate through each sheet
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)

            # Create a Word document
            doc = Document()
            doc.add_heading(f'Sheet: {sheet_name}', level=1)

            # Add table to the document
            if not df.empty:
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    hdr_cells[i].text = str(col_name)

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell)

            # Save the Word document to a byte stream
            word_io = io.BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            stream_name = f"{sheet_name}.docx"
            word_streams.append(
                    {
                        "file_id": file_id,
                        "stream": DocumentStream(name=stream_name, stream=word_io)
                    }
                )
            logger.info(f"{stream_name} Word Doc added")
        return word_streams

    async def download_file(self, site_id: str, drive_id: str, item_data: Item):
        logger.info(f"Downloading file: {item_data}")

        file_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/items/{item_data.id}/content"
        response = requests.get(file_url, headers=self._get_headers())
        if response.status_code==401:
            self.auth_service._generate_token()
            return await self.download_file(site_id=site_id,drive_id=drive_id,item_data=item_data)
        response.raise_for_status()
        document_streams = {
            "file_id": item_data.id,
            "item_object": item_data,
            "stream": DocumentStream(
                    name=item_data.path.split('/')[-1],
                    stream=io.BytesIO(response.content)
                )}
        return document_streams

    async def get_site_id(self, site_name):
        try:
            if (site_name not in self.site_mapping):
                sp_site_name = "sita365.sharepoint.com"
                site_resp = requests.get(
                    url=f"https://graph.microsoft.com/v1.0/sites/{sp_site_name}:/sites/{site_name}",
                    headers=self._get_headers())
                site_id = site_resp.json()["id"]
            else:
                sp_site_name = "sita365.sharepoint.com"
                site_id = self.site_mapping[site_name]["site_id"]

            return site_id, sp_site_name
        except Exception as e:
            logger.error(f"Failed to fetch response for site_name:{site_name} e: {str(e)}")

    async def get_drive_id(self, site_id, sp_site_name, site_name, drive_name):
        get_drives_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
        get_drives_resp = requests.get(url=get_drives_url, headers=self._get_headers())
        # print(json.dumps(get_drives_resp.json(), indent = 4))
        web_url = f"https://{sp_site_name}/sites/{site_name}/{quote(drive_name)}"
        all_drives = get_drives_resp.json()
        drive_id = None

        for drive in all_drives["value"]:
            if (drive["webUrl"] == web_url):
                drive_id = drive["id"]

        if (not drive_id):
            logger.error("Drive Not found error")

        return drive_id

    async def get_item_id(self, site_id, drive_id, file_path):
        files_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{file_path}"
        files_resp = requests.get(url=files_url, headers=self._get_headers())
        if files_resp.status_code == 200:
            item_id = files_resp.json().get('id')
            return Item(
                id=item_id,
                name=files_resp.json().get('name'),
                path=file_path,
                lastModifiedDate=get_iso_datetime_to_unix_timestamp(files_resp.json().get("lastModifiedDateTime"))
            )
        else:
            raise Exception("Unable to Fetch Item ID")

    async def get_file_spn_details(self, link):
        if '?id=' in link:
            parsed_link = urlparse(link)
            site_name = parsed_link.path.split("/")[2]
            query_string = parse_qs(parsed_link.query)
            decoded_qs = query_string["id"][0].split("/")
            drive_name = decoded_qs[3]
            file_path = "/".join(decoded_qs[4:])
        else:
            parsed_link = urlparse(link)
            url_path = unquote(parsed_link.path)
            link_after_site = url_path.split("sites/")
            site_name = link_after_site[1].split('/')[0]
            query_string = parse_qs(parsed_link.query)
            drive_name = link_after_site[1].split('/')[1]
            file_path = "/".join(link_after_site[1].split('/')[2:])

        site_id, sp_site_name = await self.get_site_id(site_name=site_name)
        drive_id = await self.get_drive_id(site_id=site_id,
                                           sp_site_name=sp_site_name,
                                           site_name=site_name,
                                           drive_name=drive_name
                                           )
        item_object = await self.get_item_id(site_id=site_id,
                                             drive_id=drive_id,
                                             file_path=file_path
                                             )

        file_object = await self.download_file(
            site_id=site_id, drive_id=drive_id, item_data=item_object
        )
        logger.info(f"File Object Fetchec for file_id:{file_object.get('file_id')}")
        return file_object

    async def download_file_independent(self, site_id: str, drive_id: str, file_id: str,file_path :str):
        logger.info(f"Downloading file: {file_id}")

        pptx_as_pdf = False  # Flag to indicate if PPTX was successfully downloaded as PDF
        
        # For PPTX files: Download as PDF to enable page image generation
        # This allows Docling's PDF pipeline to generate full slide images (page-0.png, page-1.png, etc.)
        # which are then processed by Vision API for comprehensive slide analysis
        if file_path.lower().endswith('.pptx'):
            logger.info(f"PPTX detected - attempting to download as PDF for slide-level image extraction")
            file_url = GraphUrls.file_content_pdf_url.value.format(site_id=site_id, drive_id=drive_id, file_id=file_id)
            
            try:
                response = requests.get(file_url, headers=self._get_headers(), allow_redirects=True)
                
                if response.status_code == 401:
                    self.auth_service._generate_token()
                    return await self.download_file_independent(site_id=site_id,
                                                                drive_id=drive_id,
                                                                file_id=file_id,
                                                                file_path=file_path)
                
                # If we get 400 error (Bad Request), fall back to downloading as PPTX
                if response.status_code == 400 or response.status_code == 406:
                    logger.warning(f"Failed to download PPTX as PDF (400 error) - falling back to PPTX download")
                    file_url = GraphUrls.file_content_url.value.format(site_id=site_id, drive_id=drive_id, file_id=file_id)
                    response = requests.get(file_url, headers=self._get_headers(), allow_redirects=True)
                    response.raise_for_status()
                    pptx_as_pdf = False  # Use standard flow, not hybrid
                else:
                    response.raise_for_status()
                    pptx_as_pdf = True  # Successfully downloaded as PDF - use hybrid flow
                    logger.info(f"Successfully downloaded PPTX as PDF - will use hybrid processing")
                    
            except requests.exceptions.HTTPError as e:
                logger.error(f"HTTP error downloading PPTX as PDF: {e} - falling back to PPTX download")
                file_url = GraphUrls.file_content_url.value.format(site_id=site_id, drive_id=drive_id, file_id=file_id)
                response = requests.get(file_url, headers=self._get_headers(), allow_redirects=True)
                response.raise_for_status()
                pptx_as_pdf = False  # Use standard flow, not hybrid
        else:
            # For other formats (DOCX, PDF, XLSX, etc.) - download as-is
            file_url = GraphUrls.file_content_url.value.format(site_id=site_id, drive_id=drive_id, file_id=file_id)
            response = requests.get(file_url, headers=self._get_headers(), allow_redirects=True)
            
            if response.status_code==401:
                self.auth_service._generate_token()
                return await self.download_file_independent(site_id=site_id,
                                                            drive_id=drive_id,
                                                            file_id=file_id,
                                                            file_path=file_path)
            response.raise_for_status()
        
        if ".xlsx" in file_path:
            logger.info(f"file is an excel")
            document_streams = self.convert_excel_to_word_streams(response.content,file_id)
        else:
            document_streams = {
                "file_id": file_id,
                "stream": DocumentStream(
                        name=file_path.split('/')[-1],
                        stream=io.BytesIO(response.content)
                    ),
                "pptx_as_pdf": pptx_as_pdf  # Flag to control hybrid vs standard processing
        }
        return document_streams

    async def get_file_path(self, link):
        if '?id=' in link:
            parsed_link = urlparse(link)
            site_name = parsed_link.path.split("/")[2]
            query_string = parse_qs(parsed_link.query)
            decoded_qs = query_string["id"][0].split("/")
            drive_name = decoded_qs[3]
            file_path = "/".join(decoded_qs[4:])
        else:
            parsed_link = urlparse(link)
            url_path = unquote(parsed_link.path)
            link_after_site = url_path.split("sites/")
            site_name = link_after_site[1].split('/')[0]
            query_string = parse_qs(parsed_link.query)
            drive_name = link_after_site[1].split('/')[1]
            file_path = "/".join(link_after_site[1].split('/')[2:])
        return file_path
